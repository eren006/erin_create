#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从「生成指令手册.py」里的 SECTIONS 数据生成 Word 版指令手册。
与 PDF 版共用同一份数据源（SECTIONS），改指令只需改一处，两份手册一起同步。
"""

import importlib.util
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# ── 加载 SECTIONS（源文件名含中文，用 importlib 按路径加载）──────────
_spec = importlib.util.spec_from_file_location("cmd_manual_src", os.path.join(HERE, "生成指令手册.py"))
_src = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_src)
SECTIONS = _src.SECTIONS

CN_FONT = "PingFang SC"
COLOR_ADMIN = RGBColor(0xC0, 0x39, 0x2B)
COLOR_PLAYER = RGBColor(0x1A, 0x7A, 0x3F)
COLOR_HEAD_BG = "2C3E50"
COLOR_ADMIN_BG = "FFF3E0"
COLOR_PLAYER_BG = "F0FFF4"


def set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_font(run, size=10, bold=False, color=None):
    run.font.name = CN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)


def add_title(doc, text, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=15, bold=True)
    return p


def add_cmd_table(doc, rows_data):
    headers = ["指令", "句号", "权限", "功能说明", "格式 / 示例"]
    widths = [3.4, 1.2, 1.8, 5.6, 5.6]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = Cm(widths[i])
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(hdr_cells[i], COLOR_HEAD_BG)

    for cmd_txt, dot, perm, desc, usage in rows_data:
        row = table.add_row().cells
        is_admin = "管理员" in perm
        color = COLOR_ADMIN if is_admin else COLOR_PLAYER
        bg = COLOR_ADMIN_BG if is_admin else COLOR_PLAYER_BG
        dot_sym = "。" if dot is True else ("" if dot is False else "。/")

        vals = [cmd_txt, dot_sym, perm, desc, usage.replace("\n", " / ") if usage else ""]
        for i, val in enumerate(vals):
            row[i].width = Cm(widths[i])
            p = row[i].paragraphs[0]
            if i in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_font(r, size=9.5, bold=(i == 0), color=color if i in (1, 2) else None)
            set_cell_bg(row[i], bg)

    return table


def build(out_path=None):
    out_path = out_path or os.path.join(HERE, "长日系统指令手册.docx")

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_title(doc, "长日系统 指令手册", size=24)
    add_title(doc, "Command Reference", size=12, bold=False, space_after=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("前缀说明：。= 中文句号前缀   /   无前缀 = 直接发送文字触发   /   * = 占位符")
    set_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))

    for title, rows in SECTIONS:
        add_section_heading(doc, title)
        add_cmd_table(doc, rows)

    doc.save(out_path)
    print(f"✅ docx 已生成：{out_path}")


if __name__ == "__main__":
    build()
