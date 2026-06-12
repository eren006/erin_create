#!/usr/bin/env python3
"""生成长日工作流 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面设置：A4，1 英寸边距 ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = int(11906)  # A4
section.page_height = int(16838)
m = int(914400 * 0.9)  # ~0.9 英寸
section.top_margin    = m
section.bottom_margin = m
section.left_margin   = m
section.right_margin  = m

# ── 默认字体 ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 颜色常量 ──────────────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0x1F, 0x39, 0x64)   # 深蓝
C_H1      = RGBColor(0x2E, 0x74, 0xB5)   # 蓝
C_H2      = RGBColor(0x2E, 0x74, 0xB5)   # 蓝
C_H3      = RGBColor(0x40, 0x40, 0x40)   # 深灰
C_WARN    = RGBColor(0xC0, 0x50, 0x00)   # 橙色警告
C_CODE_BG = RGBColor(0xF2, 0xF2, 0xF2)  # 浅灰背景
C_ADMIN   = RGBColor(0x70, 0x30, 0xA0)   # 紫色（管理员专用）
C_MUTED   = RGBColor(0x60, 0x60, 0x60)   # 说明文字灰

# ── helper 函数 ───────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, color=None, size=None, name=None):
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    if name:  run.font.name = name

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = color or C_H1
        # 下划线装饰
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '%02X%02X%02X' % (color or C_H1))
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = color or C_H2
    elif level == 3:
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = color or C_H3
    return p

def add_body(text, color=None, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color:  run.font.color.rgb = color
    if italic: run.italic = True
    return p

def add_note(text):
    """橙色说明行"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run('⚠  ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_WARN
    return p

def add_tip(text):
    """灰色小字说明"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_MUTED
    run.italic = True
    return p

def add_cmd(cmd, desc=None, is_admin=False):
    """指令行：等宽 + 可选说明"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(cmd)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = C_ADMIN if is_admin else RGBColor(0x1A, 0x1A, 0x4A)
    if desc:
        r2 = p.add_run('   ' + desc)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = C_MUTED
    return p

def add_step(num, title, body_lines=None, note=None):
    """编号步骤块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'Step {num}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = C_H1
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_H3
    if body_lines:
        for line in body_lines:
            add_cmd(line, is_admin=True)
    if note:
        add_tip(note)

def add_table(rows, col_widths=None):
    """简单表格：第一行为表头"""
    n_cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            if ri == 0:  # header
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '2E74B5')
                tcPr.append(shd)
            elif ri % 2 == 0:  # zebra
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EEF3F9')
                tcPr.append(shd)
    tbl.paragraph_format_after = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# 封面标题
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('长日系统  工作流手册')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = C_TITLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(18)
r2 = p2.add_run('运营操作全流程 + 管理员指令速查')
r2.font.size = Pt(11)
r2.font.color.rgb = C_MUTED
r2.italic = True

add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 一、开新季度
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('一、开新季度', 1)

add_heading('前提确认', 2)
add_body('• 上一季度已执行「结束季度」（若有）')
add_body('• rparchive 服务器已启动')
add_body('• 海豹后台已填写「RP存档服务器地址」和「RP存档Token」')

add_heading('操作步骤', 2)

add_step(1, '清空数据', ['.MASTER JSCLEAR changri'],
         note='不可跳过。角色存储不为空时「创建新季度」会拒绝执行。')

add_step(2, '配置基础设置', ['。设置 基础设置'],
         note='填写并回传：恋综名、戏群、后台群、公告群、水群。')

add_step(3, '创建季度')
add_cmd('。创建新季度 恋综名 复盘', is_admin=True)
add_tip('复盘模式：所有戏和场次存档到 rparchive。')
add_cmd('。创建新季度 恋综名 不复盘', is_admin=True)
add_tip('不复盘模式：不存档戏内容，仅记录玩家统计数据。')

add_step(4, '创建角色')
add_cmd('创建新角色 角色名', desc='玩家自行执行，或管理员统一创建')
add_cmd('创建NPC 角色名', desc='NPC 角色（复盘模式必须创建，否则无法自动复盘）', is_admin=True)

add_step(5, '检查功能配置', ['。设置 功能开关', '。设置 互动参数', '。设置 信件与礼品'],
         note='可按需调整，首次使用时默认值通常无需修改。')

add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 二、日常运营
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('二、日常运营', 1)

add_heading('推进日期', 2)
add_body('• 开启自动推进（auto_day_reset_enabled）时，系统到点自动推进，无需手动操作。')
add_body('• 未开启时，管理员手动执行推进指令。')

add_heading('发起官约', 2)
add_cmd('。发起官约 D1 14:00-16:00 咖啡馆 张三/李四', is_admin=True)
add_tip('格式：日期  时间段  地点  参与者（斜线分隔）。执行后自动创建约会群组并通知参与者。')

add_heading('监听', 2)
add_body('开启「自动监听所有群组」后，系统在检测到戏段开始时自动计时，无需手动操作。')
add_table([
    ['指令', '说明'],
    ['。提醒超时', '向所有超时场次发送催稿提醒（管理员）'],
    ['查看计时器', '查看所有群计时器当前状态（管理员）'],
    ['查看进行中', '查看正在进行的所有场次（管理员）'],
], col_widths=[7, 9.5])

add_heading('结戏', 2)
add_cmd('结束私约', desc='玩家手动结束当前场次')
add_tip('超时未结束时，系统自动结算并通知。')

add_heading('常用查看', 2)
add_table([
    ['指令', '说明', '权限'],
    ['玩家名单', '当前所有绑定角色', '全员'],
    ['角色卡', '自己的角色信息', '全员'],
    ['本场统计', '当前场次字数 / 段落', '全员'],
    ['查看进行中', '所有活跃场次', '管理员'],
    ['查看全员统计', '全员数据汇总', '管理员'],
], col_widths=[5.5, 8.5, 2.5])

add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 三、结束季度
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('三、结束季度', 1)

add_heading('前提确认', 2)
add_body('• 所有活跃场次已结束（「查看计时器」无进行中）')
add_body('• 若设置了「结戏前要求复盘」，所有玩家已完成复盘')

add_heading('操作步骤', 2)

add_step(1, '封存季度', ['。结束季度'],
         note='系统调用 rparchive end_season，成功后返回公开存档 URL。')

add_step(2, '确认存档')
add_body('访问返回的 URL，核对存档内容（角色列表、场次记录、数据统计）无误。', indent=True)

add_step(3, '清空数据', ['.MASTER JSCLEAR changri'],
         note='执行后数据不可恢复。务必在确认存档无误后操作。')

add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 四、管理员指令速查
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('四、管理员指令速查', 1, color=C_ADMIN)

add_heading('角色管理', 2)
add_table([
    ['指令', '说明'],
    ['。清除玩家 角色名', '删除角色及其所有数据'],
    ['设定 XX 为NPC', '切换角色的 NPC 身份'],
    ['。发起官约 日期 时间段 地点 A/B', '创建官方约会'],
    ['。提醒超时', '向超时场次发送提醒'],
], col_widths=[8.5, 8])

add_heading('关系线管理', 2)
add_table([
    ['指令', '说明'],
    ['。设置强制关系线 角色A 角色B 描述', '系统设定关系，不占发起名额'],
    ['。删除关系线 角色A 角色B', '删除两人关系线'],
    ['。清空关系线 MMDD', '清空整个平台关系线（需输入当日日期码，如 0526）'],
    ['关系线统计', '查看全员关系线数量'],
], col_widths=[8.5, 8])

add_heading('设置面板', 2)
add_table([
    ['指令', '说明'],
    ['。设置 基础设置', '群组配置、关系线开关、字数上限'],
    ['。设置 功能开关', '各功能模块开启 / 关闭'],
    ['。设置 互动参数', '监听超时、冷却时间等'],
    ['。设置 信件与礼品', '短信 / 礼物 / 心动信 / 心愿 / 道具'],
    ['。设置 目击设置', 'DLC：目击系统（需开启 DLC）'],
    ['。设置 复盘设置', 'DLC：复盘系统'],
    ['。设置 拍卖设置', 'DLC：拍卖系统'],
], col_widths=[8.5, 8])

add_heading('数据同步', 2)
add_table([
    ['指令', '说明'],
    ['。全量同步', '从 rparchive 拉取物品 / 属性 / 奖励配置'],
    ['。全量同步 预览', '仅预览，不写入'],
], col_widths=[8.5, 8])

add_divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 五、玩家指令速查
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('五、玩家指令速查', 1)

add_heading('角色', 2)
add_table([
    ['指令', '说明'],
    ['创建新角色 角色名', '绑定角色'],
    ['角色卡', '查看自己的角色卡'],
    ['玩家名单', '全员名单'],
    ['修改名字 新名字', '修改角色名'],
    ['修改皮相 明星名', '修改皮相（2 小时冷却）'],
    ['修改签名 内容', '修改签名（12 小时冷却）'],
    ['额外账号 QQ号', '绑定多账号'],
], col_widths=[7.5, 9])

add_heading('互动', 2)
add_table([
    ['指令', '说明'],
    ['电话 对方名 时间段', '发起电话约会'],
    ['私约 对方名 时间段 地点', '发起私密约会'],
    ['短信 对方名 内容', '短信'],
    ['送礼 对方名 礼物名', '送礼物'],
    ['结束私约', '手动结束当前场次'],
], col_widths=[7.5, 9])

add_heading('关系线', 2)
add_table([
    ['指令', '说明'],
    ['拉线 对方名 内容', '发起关系线并添加细节'],
    ['确认关系线 对方名', '确认对方发起的关系线'],
    ['查看关系线', '查看自己所有关系'],
    ['查看关系线 角色名', '查看与某角色的完整细节（合并转发）'],
    ['撤回关系 对方名 内容', '撤回某条细节（精确匹配）'],
], col_widths=[7.5, 9])

# ── 保存 ───────────────────────────────────────────────────────────────────────
out = '/Users/erinren/erin_creation/长日系统/长日工作流.docx'
doc.save(out)
print(f'已保存：{out}')
